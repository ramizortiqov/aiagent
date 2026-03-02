import asyncio
import os
import logging
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import FSInputFile
import re
# === НОВАЯ БИБЛИОТЕКА GOOGLE ===
from google import genai
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt, RGBColor,Cm 
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import PyPDF2
from docx.oxml.ns import qn

# ================= НАСТРОЙКИ =================
TELEGRAM_TOKEN = "8489524872:AAFRy5x8W15tGhcLdMKtEsKh3EmKHXRd_DM"

# === ВСТАВЬ СЮДА ВСЕ СВОИ 8 КЛЮЧЕЙ ===
API_KEYS = [
     # Ключ 2
]

MODEL_NAME = 'gemini-flash-latest'

# Глобальная переменная для переключения ключей
current_key_index = 0
TEMPLATE_FILE = "temp.docx"

# === ИНСТРУКЦИЯ ===
SYSTEM_INSTRUCTION = """
ТРЕБОВАНИЯ К ОФОРМЛЕНИЮ (СТРОГО КАК В УЧЕБНИКЕ):
1. Заголовок вопроса должен быть жирным и пронумерованным.
2. Текст должен быть плотным.
3. Используй маркированные списки (* пункт).
4. Ключевые термины выделяй жирным (**термин**).
5. Для сравнений используй таблицы.
6. Если нужны подзаголовки внутри ответа, используй ###.

ФОРМАТ ОТВЕТА:
**№. Текст вопроса**
Текст ответа...

### Подзаголовок (если нужен)
* Пункт 1
* Пункт 2

Списки пиши СЛИТНО, без пустых строк между пунктами.

| Параметр | Значение |
| --- | --- |
| Данные | Данные |

___LINE___
"""

logging.basicConfig(level=logging.INFO)
bot = Bot(token=TELEGRAM_TOKEN)
dp = Dispatcher()


def smart_join_lines(text):
    """
    Умная склейка разорванных строк.
    Если строка не является списком или заголовком, она приклеивается к предыдущей.
    """
    if not text:
        return ""

    text = re.sub(r'\n{2,}', '\n', text)

    lines = text.split('\n')
    joined_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        is_list_item = (
            stripped.startswith('* ') or
            stripped.startswith('- ') or
            stripped.startswith('• ') or
            stripped.startswith('· ')
        )
        is_header = stripped.startswith('#') or (
            len(stripped) > 2 and stripped[0].isdigit() and '.' in stripped[:4]
        )
        is_table = stripped.startswith('|')
        is_line = stripped == "___LINE___"

        if joined_lines and not (is_list_item or is_header or is_table or is_line):
            prev_line = joined_lines[-1]
            prev_is_list = (
                prev_line.startswith('• ') or
                prev_line.startswith('* ') or
                prev_line.startswith('- ')
            )
            prev_is_special = prev_line.startswith(
                ('*', '-', '#', '|', '___', '•')
            ) or (
                len(prev_line) > 2
                and prev_line[0].isdigit()
                and '.' in prev_line[:4]
            )

            # Не склеиваем элементы списка между собой
            if is_list_item and prev_is_list:
                joined_lines.append(stripped)
                continue

            if not prev_is_special:
                joined_lines[-1] = joined_lines[-1] + " " + stripped
                continue

        joined_lines.append(stripped)

    return "\n".join(joined_lines)


# === ГРАНИЦЫ ТАБЛИЦЫ ===
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


# === ФОРМАТИРОВАНИЕ ===
def apply_formatting(paragraph, text, is_list=False):
    p_fmt = paragraph.paragraph_format
    p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_fmt.contextual_spacing = False

    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)

    if is_list:
        # Для списков: висячий отступ
        p_fmt.left_indent = Cm(1.27)
        p_fmt.first_line_indent = Cm(-0.63)
        p_fmt.tab_stops.clear_all()
    else:
        # ★ НОВОЕ: Отступ первой строки 1,25 см для обычного текста
        p_fmt.first_line_indent = Cm(1.25)

    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if i % 2 == 1:
            run.bold = True

    for run in paragraph.runs:
        r = run._element
        rPr = r.get_or_add_rPr()
        if rPr.rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        else:
            rFonts = rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def create_word_table(doc, markdown_lines):
    data = []
    for line in markdown_lines:
        if not line.strip():
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        data.append(cells)

    clean_data = [
        row for row in data
        if not all(set(c).issubset({'-', ':', ' '}) for c in row)
    ]
    if not clean_data:
        return

    rows = len(clean_data)
    cols = max(len(r) for r in clean_data)

    table = doc.add_table(rows=rows, cols=cols)
    set_table_borders(table)

    for r, row_data in enumerate(clean_data):
        row_cells = table.rows[r].cells
        for c, cell_text in enumerate(row_data):
            if c < len(row_cells):
                cell_p = row_cells[c].paragraphs[0]
                cell_p.text = ""

                # ★ НОВОЕ: форматируем ячейку
                p_fmt = cell_p.paragraph_format
                p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
                p_fmt.first_line_indent = Cm(0)  # В таблице без отступа

                parts = cell_text.split('**')
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = cell_p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if i % 2 == 1:
                        run.bold = True
                    # ★ НОВОЕ: первая строка таблицы — жирная
                    if r == 0:
                        run.bold = True

                for run in cell_p.runs:
                    rEl = run._element
                    rPr = rEl.get_or_add_rPr()
                    if rPr.rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    else:
                        rFonts = rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def parse_and_add_content(doc, full_text):
    cleaned_text = smart_join_lines(full_text)
    lines = cleaned_text.split('\n')

    table_buffer = []
    is_inside_table = False

    for line in lines:
        stripped = line.strip()

        # Таблицы
        if stripped.startswith('|') and stripped.endswith('|'):
            table_buffer.append(stripped)
            is_inside_table = True
            continue
        if is_inside_table:
            create_word_table(doc, table_buffer)
            table_buffer = []
            is_inside_table = False
        if not stripped:
            continue

        # Линия
        if stripped == "___LINE___":
            p = doc.add_paragraph()
            run = p.add_run("_" * 85)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # ★ ИСПРАВЛЕНО: Списки — с маркером, БЕЗ отступа первой строки 1,25
        if (stripped.startswith('* ') or stripped.startswith('- ') or
                stripped.startswith('• ') or stripped.startswith('· ')):

            clean_text = stripped[2:].strip()

            p = doc.add_paragraph()

            # Маркер вручную
            marker_run = p.add_run('• ')
            marker_run.font.name = 'Times New Roman'
            marker_run.font.size = Pt(12)

            # Текст с форматированием (**жирный**)
            parts = clean_text.split('**')
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if i % 2 == 1:
                    run.bold = True

            # Отступы для списка (БЕЗ 1,25 см)
            p_fmt = p.paragraph_format
            p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_fmt.left_indent = Cm(1.27)
            p_fmt.first_line_indent = Cm(-0.63)
            p_fmt.space_after = Pt(0)
            p_fmt.space_before = Pt(0)
            p_fmt.tab_stops.clear_all()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for run in p.runs:
                rEl = run._element
                rPr = rEl.get_or_add_rPr()
                if rPr.rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                else:
                    rFonts = rPr.rFonts
                rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            continue

        # Заголовки (###)
        if stripped.startswith('#'):
            clean_text = stripped.lstrip('#').strip()
            p = doc.add_paragraph()
            apply_formatting(p, f"**{clean_text}**")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        # Основной текст — с отступом 1,25 см (через apply_formatting)
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
        else:
            p = doc.add_paragraph()

        apply_formatting(p, stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if table_buffer:
        create_word_table(doc, table_buffer)


# === ЛОГИКА БОТА ===
async def extract_text_from_file(file_path, file_ext):
    questions = []
    try:
        if file_ext == '.docx':
            doc = Document(file_path)
            questions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        elif file_ext == '.pdf':
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    extract = page.extract_text()
                    if extract:
                        text += extract + "\n"
            questions = [line.strip() for line in text.split('\n') if line.strip()]
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                questions = [line.strip() for line in f.readlines() if line.strip()]
    except Exception as e:
        logging.error(f"Ошибка чтения: {e}")
        return None
    return questions


async def process_with_gemini(questions_batch, start_number):
    global current_key_index
    prompt_text = f"{SYSTEM_INSTRUCTION}\n\nВот вопросы (начиная с номера {start_number}):\n"
    for i, q in enumerate(questions_batch, start_number):
        prompt_text += f"{i}. {q}\n"

    while True:
        active_key = API_KEYS[current_key_index]
        client = genai.Client(api_key=active_key)
        try:
            response = await client.aio.models.generate_content(
                model=MODEL_NAME,
                contents=prompt_text
            )
            print("=== ОТВЕТ GEMINI (СЫРОЙ) ===")
            print(response.text)
            return response.text

        except Exception as e:
            error_str = str(e).lower()
            if "429" in error_str or "quota" in error_str or "resource exhausted" in error_str:
                logging.warning(f"⚠️ Ключ {current_key_index+1} устал.")
                current_key_index = (current_key_index + 1) % len(API_KEYS)
                await asyncio.sleep(1)
                continue
            else:
                logging.error(f"Ошибка API: {e}")
                return f"Ошибка: {e}"


@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    await message.answer("Привет! Исправил списки, таблицы и отступы.")


@dp.message(F.document)
async def handle_document(message: types.Message):
    document = message.document
    file_name = document.file_name
    file_ext = os.path.splitext(file_name)[1].lower()

    if file_ext not in ['.docx', '.txt', '.pdf']:
        await message.answer("Нужен файл .docx, .pdf или .txt")
        return

    if not os.path.exists(TEMPLATE_FILE):
        await message.answer(f"⚠️ Ошибка: Нет файла {TEMPLATE_FILE}!")
        return

    wait_msg = await message.answer("Файл принят. Начинаю... ⏳")

    file_info = await bot.get_file(document.file_id)
    downloaded_file = await bot.download_file(file_info.file_path)
    temp_input = f"temp_{document.file_id}{file_ext}"
    with open(temp_input, 'wb') as f:
        f.write(downloaded_file.read())

    questions = await extract_text_from_file(temp_input, file_ext)

    if not questions:
        await message.answer("Файл пуст.")
        os.remove(temp_input)
        return

    BATCH_SIZE = 4
    all_responses_text = []

    for i in range(0, len(questions), BATCH_SIZE):
        batch = questions[i : i + BATCH_SIZE]
        start_number = i + 1

        await bot.edit_message_text(
            f"Обрабатываю вопросы {start_number}-{start_number + len(batch) - 1} из {len(questions)}...",
            chat_id=message.chat.id,
            message_id=wait_msg.message_id
        )

        answer_text = await process_with_gemini(batch, start_number)
        all_responses_text.append(answer_text)
        await asyncio.sleep(5)

    result_doc = Document(TEMPLATE_FILE)

    for block in all_responses_text:
        parse_and_add_content(result_doc, block)

    output_filename = f"Answers_{file_name}.docx"
    result_doc.save(output_filename)

    await message.answer_document(FSInputFile(output_filename), caption="Готово! ✅ Проверяй.")

    if os.path.exists(temp_input):
        os.remove(temp_input)
    if os.path.exists(output_filename):
        os.remove(output_filename)


async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
